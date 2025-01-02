import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
import json

def load_members():
    try:
        with open("members.json", "r") as file:
            return json.load(file)
    except FileNotFoundError:
        return []

def save_members():
    with open("members.json", "w") as file:
        json.dump(members, file, indent=4)

def set_font(run, font_name="Arial", font_size=12, font_color=(0, 0, 0), bold=False):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(*font_color)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

def generate_docx(data, signer, task_details, output_path):
    template = Document("SPT_TEMPLATE.docx")

    for paragraph in template.paragraphs:
        if "{date}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{date}", datetime.now().strftime("%d %B %Y"))
        for run in paragraph.runs:
            set_font(run) 
    tables = template.tables

    # Table 1: Yang Bertanda Tangan
    header_table = tables[0]
    header_data = [
        signer['name'],
        signer['nip'],
        signer['pangkat'],
        signer['jabatan'],
        signer['organization']
    ]

    for i, value in enumerate(header_data):
        if len(header_table.rows[i].cells) > 1:  
            cell = header_table.cell(i, 2)
            cell.text = value
            for run in cell.paragraphs[0].runs:
                set_font(run, bold=(i == 0))  

    # Table 2: Yang Bertugas
    assignments_table = tables[1]
    current_row_idx = 0

    required_columns = 4
    if len(assignments_table.rows[0].cells) < required_columns:
        messagebox.showerror(
            "Error",
            "The table must have at least 6 columns.",
        )
        return
    
    field_names_translated = ["Nama", "NIP", "Pangkat/Golongan", "Jabatan", "Satuan Organisasi"]
    for row_idx, member in enumerate(data):
        for field_idx, field_name in enumerate(["name", "nip", "pangkat", "jabatan", "organization"]):
            if current_row_idx < len(assignments_table.rows):
                row = assignments_table.rows[current_row_idx].cells
            else:
                row = assignments_table.add_row().cells
            
            row[0].text = str(row_idx + 1) if field_idx == 0 else ""
            row[1].text = field_names_translated[field_idx]
            row[2].text = ":"
            row[3].text = member[field_name]
            paragraph1 = row[0].paragraphs[0]
            paragraph1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph2 = row[2].paragraphs[0]
            paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
 
            for cell in row:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = Pt(12)  # Line spacing
                    paragraph.paragraph_format.space_before = Pt(0)  # Space before paragraph
                    paragraph.paragraph_format.space_after = Pt(0)  # Space after paragraph
                for run in paragraph.runs:
                    is_name_field = field_name == "name" and cell is row[3]
                    set_font(run, bold=is_name_field)

            current_row_idx += 1

        if row_idx < len(data) - 1:
            if current_row_idx < len(assignments_table.rows):
                row = assignments_table.rows[current_row_idx].cells
            else:
                row = assignments_table.add_row().cells
            for cell in row:
                cell.text = ""
            current_row_idx += 1

    # Table 3: Detail Tugas
    task_table = tables[2]
    task_table.cell(0, 2).text = task_details["tugas"]
    task_table.cell(1, 2).text = task_details["lama_perjalanan"]
    task_table.cell(2, 2).text = task_details["lokasi"]
    task_table.cell(3, 2).text = task_details["tanggal_berangkat"]
    task_table.cell(4, 2).text = task_details["sumber_dana"]

    for row in task_table.rows:
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                set_font(run)

    # Table 4: Tanda Tangan
    footer_table = tables[3]
    current_month_year = datetime.now().strftime("%B %Y")
    footer_table.cell(0, 0).text = f"Jakarta,    {current_month_year}"
    footer_table.cell(1, 0).text = signer['jabatan']
    footer_table.cell(3, 0).text = signer['name']

    for row in footer_table.rows:
        for cell in row.cells:
            for run in cell.paragraphs[0].runs:
                set_font(run)

    template.save(output_path)
def open_add_members_window():
    def add_members():
        name = name_entry.get().strip()
        nip = nip_entry.get().strip()
        pangkat = pangkat_entry.get().strip()
        jabatan = jabatan_entry.get().strip()
        organization = organization_entry.get().strip()

        if not (name and nip and pangkat and jabatan and organization):
            messagebox.showerror("Error", "All fields are required!")
            return

        new_member = {
            "name": name,
            "nip": nip,
            "pangkat": pangkat,
            "jabatan": jabatan,
            "organization": organization,
        }
        members.append(new_member)
        members_list.insert("end", f"{name} ({nip})")

        save_members()
        add_window.destroy()

    
    add_window = tk.Toplevel(root)
    add_window.title("Tambah Anggota")
    add_window.geometry("400x300")

    frame = ttk.Frame(add_window)
    frame.grid(pady=10, padx=10)

    ttk.Label(frame, text="Nama:").grid(row=0, column=0, sticky="w")
    name_entry = ttk.Entry(frame, width=40)
    name_entry.grid(row=0, column=1, pady=5)

    ttk.Label(frame, text="NIP:").grid(row=1, column=0, sticky="w")
    nip_entry = ttk.Entry(frame, width=40)
    nip_entry.grid(row=1, column=1, pady=5)

    ttk.Label(frame, text="Pangkat:").grid(row=2, column=0, sticky="w")
    pangkat_entry = ttk.Entry(frame, width=40)
    pangkat_entry.grid(row=2, column=1, pady=5)

    ttk.Label(frame, text="Jabatan:").grid(row=3, column=0, sticky="w")
    jabatan_entry = ttk.Entry(frame, width=40)
    jabatan_entry.grid(row=3, column=1, pady=5)

    ttk.Label(frame, text="Satuan Kerja:").grid(row=4, column=0, sticky="w")
    organization_entry = ttk.Entry(frame, width=40)
    organization_entry.grid(row=4, column=1, pady=5)

    add_button = ttk.Button(frame, text="Tambah", command=add_members)
    add_button.grid(row=5, column=0, columnspan=2, pady=10)

def edit_members():
    selected_indices = members_list.curselection()
    if not selected_indices or len(selected_indices) > 1:
        messagebox.showerror("Error", "Please select exactly one member to edit!")
        return

    index = selected_indices[0]
    member = members[index]

    def update_member():
        name = name_entry.get().strip()
        nip = nip_entry.get().strip()
        pangkat = pangkat_entry.get().strip()
        jabatan = jabatan_entry.get().strip()
        organization = organization_entry.get().strip()

        if not (name and nip and pangkat and jabatan and organization):
            messagebox.showerror("Error", "All fields are required!")
            return

      
        members[index] = {
            "name": name,
            "nip": nip,
            "pangkat": pangkat,
            "jabatan": jabatan,
            "organization": organization,
        }

        members_list.delete(index)
        members_list.insert(index, f"{name} ({nip})")

        save_members()
        edit_window.destroy()

    edit_window = tk.Toplevel(root)
    edit_window.title("Edit Member")
    edit_window.geometry("400x300")

    frame = ttk.Frame(edit_window)
    frame.grid(pady=10, padx=10)

    ttk.Label(frame, text="Nama:").grid(row=0, column=0, sticky="w")
    name_entry = ttk.Entry(frame, width=40)
    name_entry.insert(0, member["name"])
    name_entry.grid(row=0, column=1, pady=5)

    ttk.Label(frame, text="NIP:").grid(row=1, column=0, sticky="w")
    nip_entry = ttk.Entry(frame, width=40)
    nip_entry.insert(0, member["nip"])
    nip_entry.grid(row=1, column=1, pady=5)

    ttk.Label(frame, text="Pangkat:").grid(row=2, column=0, sticky="w")
    pangkat_entry = ttk.Entry(frame, width=40)
    pangkat_entry.insert(0, member["pangkat"])
    pangkat_entry.grid(row=2, column=1, pady=5)

    ttk.Label(frame, text="Jabatan:").grid(row=3, column=0, sticky="w")
    jabatan_entry = ttk.Entry(frame, width=40)
    jabatan_entry.insert(0, member["jabatan"])
    jabatan_entry.grid(row=3, column=1, pady=5)

    ttk.Label(frame, text="Satuan Kerja:").grid(row=4, column=0, sticky="w")
    organization_entry = ttk.Entry(frame, width=40)
    organization_entry.insert(0, member["organization"])
    organization_entry.grid(row=4, column=1, pady=5)

    update_button = ttk.Button(frame, text="Update", command=update_member)
    update_button.grid(row=5, column=0, columnspan=2, pady=10)

def delete_members():
    selected_indices = members_list.curselection()
    if not selected_indices:
        messagebox.showerror("Error", "No members selected for deletion!")
        return

    for idx in reversed(selected_indices):
        members_list.delete(idx)
        del members[idx]

    save_members()
    messagebox.showinfo("Success", "Selected members have been deleted.")

def save_doc():
    selected_members = [members[idx] for idx in members_list.curselection()]
    if not selected_members:
        messagebox.showerror("Error", "No members selected!")
        return

    selected_signer = signer_dropdown.get()
    if not selected_signer:
        messagebox.showerror("Error", "No signer selected!")
        return

    signer = next((member for member in members if member["name"] == selected_signer), None)
    if not signer:
        messagebox.showerror("Error", "Selected signer not found in members list!")
        return

    # Collect task details
    task_details = {
        "tugas": task_entry.get(),
        "lama_perjalanan": duration_entry.get(),
        "lokasi": location_entry.get(),
        "tanggal_berangkat": departure_date_entry.get(),
        "sumber_dana": funding_entry.get(),
    }

    for key, value in task_details.items():
        if not value:
            messagebox.showerror("Error", f"Task detail '{key}' is required!")
            return

    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx", filetypes=[("Word Documents", "*.docx")]
    )
    if file_path:
        generate_docx(selected_members, signer, task_details, file_path)
        messagebox.showinfo("Success", f"Document saved at {file_path}")

root = tk.Tk()
root.title("Surat Tugas Generator")

members = load_members()

# 1. "Yang Bertanda Tangan" Section
ttk.Label(root, text="Yang Bertanda Tangan:").grid(row=0, column=0, padx=10, pady=5, sticky="w")
signer_dropdown = ttk.Combobox(root, values=[member["name"] for member in members], state="readonly", width=50)
signer_dropdown.grid(row=1, column=0, padx=10, pady=5, sticky="w")

# 2. "Detail Tugas" Section
ttk.Label(root, text="Detail Tugas:").grid(row=2, column=0, padx=10, pady=5, sticky="w")
task_frame = ttk.Frame(root)
task_frame.grid(row=3, column=0, padx=10, pady=5, sticky="w")

ttk.Label(task_frame, text="Tugas:").grid(row=0, column=0, sticky="w")
task_entry = ttk.Entry(task_frame, width=50)
task_entry.grid(row=0, column=1, padx=5, pady=5)

ttk.Label(task_frame, text="Lama Perjalanan:").grid(row=1, column=0, sticky="w")
duration_entry = ttk.Entry(task_frame, width=50)
duration_entry.grid(row=1, column=1, padx=5, pady=5)

ttk.Label(task_frame, text="Lokasi:").grid(row=2, column=0, sticky="w")
location_entry = ttk.Entry(task_frame, width=50)
location_entry.grid(row=2, column=1, padx=5, pady=5)

ttk.Label(task_frame, text="Tanggal Keberangkatan:").grid(row=3, column=0, sticky="w")
departure_date_entry = ttk.Entry(task_frame, width=50)
departure_date_entry.grid(row=3, column=1, padx=5, pady=5)

ttk.Label(task_frame, text="Sumber Dana:").grid(row=4, column=0, sticky="w")
funding_entry = ttk.Entry(task_frame, width=50)
funding_entry.grid(row=4, column=1, padx=5, pady=5)

# 3. List of Members Section
ttk.Label(root, text="List Anggota:").grid(row=4, column=0, padx=10, pady=5, sticky="w")
members_list = tk.Listbox(root, height=10, width=70, selectmode="multiple")
members_list.grid(row=5, column=0, padx=10, pady=5)

# Populate listbox with loaded members
for member in members:
    members_list.insert("end", f"{member['name']} ({member['nip']})")

# 4. "Tambah Anggota" Button
tambah_anggota_button = ttk.Button(root, text="Tambah Anggota", command=open_add_members_window)
tambah_anggota_button.grid(row=6, column=0, padx=10, pady=10, sticky="w")

# 5. "Delete Members" Button
delete_members_button = ttk.Button(root, text="Hapus Anggota", command=delete_members)
delete_members_button.grid(row=6, column=1, padx=10, pady=10, sticky="w")

edit_members_button = ttk.Button(root, text="Edit Member", command=edit_members)
edit_members_button.grid(row=7, column=0, padx=10, pady=10, sticky="w")

# Add search bar to UI
# search_frame = ttk.Frame(root)
# search_frame.grid(row=6, column=0, padx=10, pady=5, sticky="w")

# ttk.Label(search_frame, text="Search Members:").grid(row=0, column=0, sticky="w")
# search_entry = ttk.Entry(search_frame, width=40)
# search_entry.grid(row=0, column=1, padx=5)

# search_button = ttk.Button(search_frame, text="Search", command=search_members)
# search_button.grid(row=0, column=2, padx=5)

# 6. Save Button
save_button = ttk.Button(root, text="Save Document", command=save_doc)
save_button.grid(row=8, column=0, padx=10, pady=10, sticky="w")

# Run the application
root.mainloop()
